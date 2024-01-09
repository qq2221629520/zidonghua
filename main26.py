#将电梯位置填入限速器报告中

import os
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

data = {'1.docx': '1#楼南梯', '10.docx': '13#楼', '11.docx': '1#楼北梯', '12.docx': '10#楼1单元东梯', '13.docx': '10#楼2单元东梯', '14.docx': '10#楼2单元西梯', '15.docx': '11#楼', '16.docx': '4-2-西梯', '17.docx': '4-1-西梯', '18.docx': '4-2-东梯', '19.docx': '4-1-东梯', '2.docx': '10#楼1单元西梯', '20.docx': '7-2', '21.docx': '7-1', '22.docx': '5-2', '23.docx': '5-1', '24.docx': '6-1', '25.docx': '6-2', '26.docx': '14#楼', '27.docx': '15#楼', '28.docx': '16#楼', '3.docx': '2#楼2单元东梯', '4.docx': '2#楼2单元西梯', '5.docx': '2#楼1单元东梯', '6.docx': '2#楼1单元西梯', '7.docx': '9#楼西梯', '8.docx': '9#楼东梯', '9.docx': '12#楼'}

folder_path = r'D:\个人资料\电梯\电梯限速器动作速度校检报告 山东润华物业管理有限公司潍坊分公司 28个\电梯限速器动作速度校检报告'

# 设置样式
font_style = '宋体'
font_size = Pt(12)
alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

for filename in os.listdir(folder_path):
    if filename.endswith(".docx") and filename in data:
        file_path = os.path.join(folder_path, filename)
        doc = Document(file_path)
        table = doc.tables[0]
        
        # 获取单元格
        cell_0_6 = table.cell(1, 6)
        
        # 清空单元格内容
        cell_0_6.text = ""
        
        # 设置样式
        run_0_6 = cell_0_6.paragraphs[0].add_run(f"{data[filename]}")
        font = run_0_6.font
        font.name = font_style
        font.size = font_size
        cell_0_6.paragraphs[0].alignment = alignment
        
        doc.save(file_path)
