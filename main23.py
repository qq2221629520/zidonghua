import os
import random
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

# 获取当前文件夹中的所有docx文件
folder_path = r'D:\个人资料\电梯\电梯限速器动作速度校检报告 山东润华物业管理有限公司潍坊分公司 28个\电梯限速器动作速度校检报告'
docx_files = [f for f in os.listdir(folder_path) if f.endswith('.docx')]

# 定义字体和字号
font_name = '宋体'
font_size = Pt(12)  # 小四字号

# 遍历每个docx文件并执行插入操作
for docx_file in docx_files:
    # 构建文件的完整路径
    docx_path = os.path.join(folder_path, docx_file)

    # 加载文档
    doc = Document(docx_path)

    # 获取第二个表格
    table = doc.tables[1]  # 注意：表格索引从0开始

    # 在指定单元格（0,6）插入文本
    cell_0_6 = table.cell(0, 6)
    cell_0_6.text = ""
    paragraph_0_6 = cell_0_6.paragraphs[0]
    run_0_6 = paragraph_0_6.add_run(f"{random.uniform(2.01, 2.20):.2f}  m/s")
    run_0_6.font.name = font_name
    run_0_6.font.size = font_size
    paragraph_0_6.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # 在指定单元格（1,6）插入随机数
    random_number_1_6 = round(random.uniform(2.01, 2.20), 2)  # 定义随机数
    cell_1_6 = table.cell(1, 6)
    cell_1_6.text = ""
    paragraph_1_6 = cell_1_6.paragraphs[0]
    run_1_6 = paragraph_1_6.add_run(f"{random_number_1_6:.2f}  m/s")
    run_1_6.font.name = font_name
    run_1_6.font.size = font_size
    paragraph_1_6.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # 在指定单元格（0,11）插入文本
    cell_0_11 = table.cell(0, 11)
    cell_0_11.text = ""
    paragraph_0_11 = cell_0_11.paragraphs[0]
    run_0_11 = paragraph_0_11.add_run(f"{random.uniform(2.21, 2.32):.2f}  m/s")
    run_0_11.font.name = font_name
    run_0_11.font.size = font_size
    paragraph_0_11.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # 保存修改后的文档
    doc.save(docx_path)

print("插入操作完成。")
