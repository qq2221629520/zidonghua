#将电梯编号填入限速器报告中

import os
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

data = {'1.docx': '30103707252015030011', '10.docx': '30103707252015030020', '11.docx': '30103707252015030021', '12.docx': '30103707252015030022', '13.docx': '30103707252015030023', '14.docx': '30103707252015030024', '15.docx': '30103707252015030025', '16.docx': '30103707252016060007', '17.docx': '30103707252016060008', '18.docx': '30103707252016060009', '19.docx': '30103707252016060010', '2.docx': '30103707252015030012', '20.docx': '30103707252016060011', '21.docx': '30103707252016060012', '22.docx': '30103707252016060013', '23.docx': '30103707252016060014', '24.docx': '30103707252016060015', '25.docx': '30103707252016060016', '26.docx': '30103707252016060017', '27.docx': '30103707252016060018', '28.docx': '30103707252016060019', '3.docx': '30103707252015030013', '4.docx': '30103707252015030014', '5.docx': '30103707252015030015', '6.docx': '30103707252015030016', '7.docx': '30103707252015030017', '8.docx': '30103707252015030018', '9.docx': '30103707252015030019'}

folder_path = r'D:\个人资料\电梯\电梯限速器动作速度校检报告 山东润华物业管理有限公司潍坊分公司 28个\电梯限速器动作速度校检报告'

# 遍历文件夹中的所有Word文档,将字典里的值替换到word文档中的第一个表格的第一行第六列
for filename in os.listdir(folder_path):
    if filename.endswith(".docx"):
        file_path = os.path.join(folder_path, filename)
        doc = Document(file_path)
        table = doc.tables[0]
        cell_0_6 = table.cell(0, 6)
        cell_0_6.text = ""
        paragraph_0_6 = cell_0_6.paragraphs[0]
        run_0_6 = paragraph_0_6.add_run(f"{data[filename]}")

        # Set font, size, and alignment for cell_0_6
        run_0_6.font.name = '宋体'  # Set font to 宋体
        run_0_6.font.size = Pt(12)  # Set font size to 小四
        cell_0_6.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER  # Set horizontal alignment to center

        doc.save(file_path)
