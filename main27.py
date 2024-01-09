from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
import os

folder_path = r'D:\个人资料\电梯\电梯限速器动作速度校检报告 山东润华物业管理有限公司潍坊分公司 28个\电梯限速器动作速度校检报告'

target_text_to_replace = "山东润华物业管理股份有限公司"
replacement_text = "山东安抚机电设备有限公司"

# Iterate through each Word document in the folder
for filename in os.listdir(folder_path):
    if filename.endswith(".docx"):
        file_path = os.path.join(folder_path, filename)
        
        # Open the Word document
        doc = Document(file_path)

        # Iterate through each table in the document
        for table in doc.tables:
            # Iterate through each cell in the table
            for row in table.rows:
                for cell in row.cells:
                    # Check if the cell contains the target text to replace
                    if target_text_to_replace in cell.text:
                        # Replace the target text with the replacement text
                        cell.text = cell.text.replace(target_text_to_replace, replacement_text)

                        # Add formatting for the cell content
                        paragraph = cell.paragraphs[0]
                        run = paragraph.runs[0]
                        run.font.name = '宋体'  # Set font to 宋体
                        run.font.size = Pt(12)  # Set font size to 小四
                        paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER  # Set horizontal alignment to center

        # Save the modified document
        doc.save(file_path)
