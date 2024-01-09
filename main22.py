# 获取电梯自检报告中的20位数字 号码，储存在字典中

import os
from docx import Document
import re

# 指定文件夹路径
folder_path = r'D:\个人资料\电梯\28台\2020领仕郡28台自检报告及报检申请表\2020领仕郡28台自检报告及报检申请表'

# 遍历文件夹中的所有Word文档
for filename in os.listdir(folder_path):
    if filename.endswith(".docx"):
        file_path = os.path.join(folder_path, filename)
        
        # 使用python-docx读取文档内容
        doc = Document(file_path)
        
        # 在文档中查找连续的20位数字
        for paragraph in doc.paragraphs:
            text = paragraph.text
            # 使用正则表达式匹配20位数字
            matches = re.findall(r'\d{20}', text)
            for match in matches:
                print(f"在文件 {filename} 中找到连续的20位数字: {match}")


# 将数据储存在字典中 按照键值对的形式，文件名为键，文件中的数字为值
data={}
for filename in os.listdir(folder_path):
    if filename.endswith(".docx"):
        file_path = os.path.join(folder_path, filename)
        doc = Document(file_path)
        for paragraph in doc.paragraphs:
            text = paragraph.text
            matches = re.findall(r'\d{20}', text)
            for match in matches:
                data[filename]=match
print(data)    