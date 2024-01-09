#查找word里的表格，某个单元格里的文字，然后设置字体，字号，居中对齐

from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
import os

folder_path = r'D:\个人资料\电梯\28台\28台限速器报告\28台限速器报告\临时'

target_text = "技术参数"

# Iterate through each Word document in the folder
for filename in os.listdir(folder_path):
    if filename.endswith(".docx"):
        file_path = os.path.join(folder_path, filename)
        
        # Open the Word document
        doc = Document(file_path)

        # Iterate through each table in the document
        for table in doc.tables:
            # Iterate through each cell in the table
            for row in table.rows:
                for cell in row.cells:
                    # Check if the cell contains the target text
                    if target_text in cell.text:
                        # Set font, size, and alignment for the cell content
                        paragraph = cell.paragraphs[0]
                        run = paragraph.runs[0]
                        run.font.name = '宋体'  # Set font to 宋体
                        run.font.size = Pt(12)  # Set font size to 小四
                        paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER  # Set horizontal alignment to center

        # Save the modified document
        doc.save(file_path)
